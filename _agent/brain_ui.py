"""
Prozessia Second Brain – Chat UI
Starten: bash _agent/start_brain_ui.sh
"""

import io
import json
import os
import re
import shutil
import subprocess
import sys
import threading
from collections import defaultdict
from datetime import datetime, time, timedelta
from pathlib import Path

import anthropic
import streamlit as st

VAULT          = Path.home() / "Documents" / "Prozessia-Brain"
INBOX          = VAULT / "_inbox"
API_KEY        = os.environ.get("ANTHROPIC_API_KEY")
INDEX_PATH     = VAULT / "_agent" / "vault.index"
META_PATH      = VAULT / "_agent" / "vault_metadata.json"
MEMORY_PATH    = VAULT / "_agent" / "memory.md"
CONTEXT_PATH   = VAULT / "_agent" / "context.md"
PROZESSIA_PATH = VAULT / "_agent" / "prozessia.md"

st.set_page_config(
    page_title="Prozessia Brain",
    page_icon=None,
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Design-Token (Prozessia Brand) ──────────────────────────────────────────
# Farben direkt aus Prozessia Messe-Flyer abgeleitet:
# C_DARK #111111 · C_PURPLE #534AB7 · C_PURPLE_LIGHT #B088FF
# C_DARK_LINE #2A2A3A · C_BODY_DARK #C8C7D4 · C_GRAY_TEXT #B8B8C2

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Sans:ital,opsz,wght@0,9..40,300;0,9..40,400;0,9..40,500;0,9..40,600;1,9..40,300&family=JetBrains+Mono:wght@300;400;500&display=swap');

/* ── Variables ─────────────────────────────────────────────────────────── */
:root {
    --color-bg:             #0D0D0D;
    --color-surface:        #111111;
    --color-surface-raised: #161616;
    --color-border:         #1E1E1E;
    --color-border-subtle:  #2A2A2A;
    --color-purple:         #534AB7;
    --color-purple-light:   #B088FF;
    --color-purple-dim:     #2A2548;
    --color-text:           #FFFFFF;
    --color-text-muted:     #B8B8C2;
    --color-text-faint:     #5A5A66;
    --color-success:        #1A7A4A;
    --color-warning:        #B7660A;
    --color-error:          #C0392B;
    --font-display:         'DM Serif Display', Georgia, serif;
    --font-body:            'DM Sans', system-ui, sans-serif;
    --font-mono:            'JetBrains Mono', monospace;
}

/* ── Base ─────────────────────────────────────────────────────────────── */
*, *::before, *::after { box-sizing: border-box !important; }

html, body,
.stApp,
[data-testid="stApp"],
[data-testid="stAppViewContainer"],
[data-testid="stMain"],
[data-testid="stHeader"],
[data-testid="stBottom"],
section.main, .main {
    background-color: var(--color-bg) !important;
    color: var(--color-text-muted) !important;
    font-family: var(--font-body) !important;
    -webkit-font-smoothing: antialiased !important;
}

[data-testid="stMainBlockContainer"],
[data-testid="stVerticalBlock"],
[data-testid="stVerticalBlockBorderWrapper"],
[data-testid="element-container"],
[data-testid="stElementContainer"],
[data-testid="stHorizontalBlock"] > div,
[data-testid="stColumn"] {
    background: transparent !important;
}

.main .block-container,
[data-testid="stMainBlockContainer"] {
    padding: 0 2.5rem 4rem !important;
    max-width: 100% !important;
}

/* ── Header ───────────────────────────────────────────────────────────── */
.pb-header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 1.5rem 0 1.5rem;
    border-bottom: 1px solid var(--color-border);
    margin-bottom: 2rem;
    position: relative;
    overflow: hidden;
}

.pb-star {
    position: absolute;
    width: 2px;
    height: 2px;
    border-radius: 50%;
    background: var(--color-purple-light);
    opacity: 0.15;
    pointer-events: none;
}

.pb-wordmark {
    font-family: var(--font-display);
    font-size: 22px;
    font-weight: 400;
    letter-spacing: -0.01em;
    color: var(--color-text);
    line-height: 1.1;
}

.pb-wordmark .accent {
    font-style: italic;
    color: var(--color-purple-light);
}

.pb-meta {
    font-family: var(--font-mono);
    font-size: 9px;
    font-weight: 400;
    color: var(--color-text-faint);
    letter-spacing: 0.12em;
    text-transform: uppercase;
    margin-top: 5px;
}

.pb-chips {
    display: flex;
    gap: 5px;
    align-items: center;
}

.pb-chip {
    font-family: var(--font-mono);
    font-size: 9px;
    font-weight: 500;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    padding: 3px 8px;
    border-radius: 2px;
}

.chip-on  { background: var(--color-purple-dim); color: var(--color-purple-light); border: 1px solid var(--color-purple); }
.chip-off { background: transparent; color: var(--color-text-faint); border: 1px solid var(--color-border-subtle); }
.chip-ok  { background: rgba(26,122,74,0.12); color: #52c97a; border: 1px solid rgba(26,122,74,0.35); }

/* ── Tabs ─────────────────────────────────────────────────────────────── */
.stTabs [data-baseweb="tab-list"] {
    background: transparent !important;
    border-bottom: 1px solid var(--color-border) !important;
    gap: 0 !important;
    padding: 0 !important;
}

.stTabs [data-baseweb="tab"] {
    background: transparent !important;
    border: none !important;
    border-bottom: 2px solid transparent !important;
    padding: 12px 24px !important;
    margin-bottom: -1px !important;
    font-family: var(--font-mono) !important;
    font-size: 9px !important;
    font-weight: 500 !important;
    letter-spacing: 0.14em !important;
    text-transform: uppercase !important;
    color: var(--color-text-faint) !important;
    transition: color 0.15s ease !important;
}

.stTabs [data-baseweb="tab"]:hover {
    color: var(--color-text-muted) !important;
    background: transparent !important;
}

.stTabs [aria-selected="true"] {
    color: var(--color-text) !important;
    border-bottom-color: var(--color-purple) !important;
}

.stTabs [data-baseweb="tab-highlight"] { display: none !important; }

.stTabs [data-baseweb="tab-panel"] {
    background: transparent !important;
    padding-top: 2rem !important;
}

/* ── Section labels ───────────────────────────────────────────────────── */
.pb-label {
    font-family: var(--font-mono);
    font-size: 9px;
    font-weight: 500;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: var(--color-text-faint);
    margin: 1.75rem 0 0.75rem;
    padding-bottom: 0.5rem;
    border-bottom: 1px solid var(--color-border);
}

/* ── Buttons ──────────────────────────────────────────────────────────── */
.stButton > button {
    background: transparent !important;
    color: var(--color-text-muted) !important;
    border: 1px solid var(--color-border-subtle) !important;
    border-radius: 0 !important;
    font-family: var(--font-body) !important;
    font-size: 12px !important;
    font-weight: 400 !important;
    padding: 7px 16px !important;
    transition: all 0.15s ease !important;
    letter-spacing: 0.02em !important;
    cursor: pointer !important;
}

.stButton > button:hover {
    background: var(--color-surface-raised) !important;
    border-color: var(--color-purple) !important;
    color: var(--color-text) !important;
}

.stButton > button:focus {
    box-shadow: none !important;
    border-color: var(--color-purple) !important;
    outline: none !important;
}

.stButton > button[kind="primary"] {
    background: var(--color-purple) !important;
    border-color: var(--color-purple) !important;
    color: #fff !important;
    font-weight: 500 !important;
}

.stButton > button[kind="primary"]:hover {
    background: #6358c8 !important;
    border-color: #6358c8 !important;
}

/* ── Inputs ───────────────────────────────────────────────────────────── */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea {
    background: var(--color-surface-raised) !important;
    color: var(--color-text) !important;
    border: 1px solid var(--color-border-subtle) !important;
    border-radius: 0 !important;
    font-family: var(--font-body) !important;
    font-size: 13px !important;
}

.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: var(--color-purple) !important;
    box-shadow: none !important;
    outline: none !important;
}

.stTextInput > div > div > input::placeholder,
.stTextArea > div > div > textarea::placeholder {
    color: var(--color-text-faint) !important;
}

/* Selectbox */
.stSelectbox > div > div {
    background: var(--color-surface-raised) !important;
    color: var(--color-text-muted) !important;
    border: 1px solid var(--color-border-subtle) !important;
    border-radius: 0 !important;
    font-family: var(--font-body) !important;
}

/* Date / Time */
.stDateInput input,
.stTimeInput input {
    background: var(--color-surface-raised) !important;
    color: var(--color-text-muted) !important;
    border: 1px solid var(--color-border-subtle) !important;
    border-radius: 0 !important;
}

/* Field labels */
.stTextInput label, .stTextArea label, .stSelectbox label,
.stDateInput label, .stTimeInput label, .stFileUploader label {
    color: var(--color-text-faint) !important;
    font-family: var(--font-mono) !important;
    font-size: 9px !important;
    font-weight: 500 !important;
    letter-spacing: 0.12em !important;
    text-transform: uppercase !important;
}

/* Toggle/Checkbox */
.stToggle label span, .stCheckbox label span {
    color: var(--color-text-muted) !important;
    font-family: var(--font-body) !important;
    font-size: 12px !important;
}

/* ── Chat ─────────────────────────────────────────────────────────────── */
[data-testid="stChatInput"] > div {
    background: var(--color-surface-raised) !important;
    border: 1px solid var(--color-border-subtle) !important;
    border-radius: 0 !important;
}

[data-testid="stChatInput"] > div:focus-within {
    border-color: var(--color-purple) !important;
    box-shadow: none !important;
}

[data-testid="stChatInput"] textarea {
    color: var(--color-text) !important;
    font-size: 14px !important;
    font-family: var(--font-body) !important;
}

[data-testid="stChatInput"] textarea::placeholder {
    color: var(--color-text-faint) !important;
}

[data-testid="stChatInput"] button {
    background: var(--color-purple) !important;
    border-radius: 0 !important;
}

[data-testid="stChatMessage"] {
    background: transparent !important;
    padding: 0.4rem 0 !important;
}

.stChatMessage p {
    color: var(--color-text-muted) !important;
    font-family: var(--font-body) !important;
    font-size: 14px !important;
    line-height: 1.75 !important;
}

/* ── Cards ────────────────────────────────────────────────────────────── */
.pb-card {
    background: var(--color-surface);
    border: 1px solid var(--color-border);
    border-radius: 0;
    padding: 12px 14px;
    margin-bottom: 1px;
    transition: border-color 0.15s, background 0.15s;
    cursor: pointer;
}

.pb-card:hover {
    background: var(--color-surface-raised);
    border-color: var(--color-border-subtle);
}

.pb-card-sender {
    font-family: var(--font-body);
    font-size: 13px;
    font-weight: 600;
    color: var(--color-text);
    display: flex;
    align-items: center;
    gap: 7px;
}

.pb-card-subject {
    font-family: var(--font-body);
    font-size: 12px;
    color: var(--color-text-faint);
    margin-top: 2px;
    overflow: hidden;
    text-overflow: ellipsis;
    white-space: nowrap;
}

.pb-card-subject.unread-subj { color: var(--color-text-muted); }

.pb-card-time {
    font-family: var(--font-mono);
    font-size: 10px;
    color: var(--color-text-faint);
    font-variant-numeric: tabular-nums;
}

.pb-unread-dot {
    width: 5px;
    height: 5px;
    border-radius: 50%;
    background: var(--color-purple-light);
    flex-shrink: 0;
}

/* ── Calendar ─────────────────────────────────────────────────────────── */
.pb-day {
    font-family: var(--font-mono);
    font-size: 9px;
    font-weight: 500;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: var(--color-text-faint);
    padding: 1.25rem 0 0.6rem;
    border-bottom: 1px solid var(--color-border);
    margin-bottom: 0.75rem;
}

.pb-event {
    background: var(--color-surface);
    border: 1px solid var(--color-border);
    border-left: 2px solid var(--color-purple);
    border-radius: 0;
    padding: 10px 14px;
    margin-bottom: 1px;
    transition: border-color 0.15s, background 0.15s;
    cursor: pointer;
}

.pb-event:hover {
    background: var(--color-surface-raised);
    border-left-color: var(--color-purple-light);
}

.pb-event-time {
    font-family: var(--font-mono);
    font-size: 9px;
    font-weight: 500;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: var(--color-purple-light);
}

.pb-event-title {
    font-family: var(--font-body);
    font-size: 14px;
    font-weight: 500;
    color: var(--color-text);
    margin-top: 3px;
}

.pb-event-meta {
    font-family: var(--font-body);
    font-size: 11px;
    color: var(--color-text-faint);
    margin-top: 4px;
}

/* ── Miscellaneous ────────────────────────────────────────────────────── */
hr, .stDivider { border-color: var(--color-border) !important; margin: 1rem 0 !important; }

.stCaption, [data-testid="stCaptionContainer"] p {
    color: var(--color-text-faint) !important;
    font-family: var(--font-mono) !important;
    font-size: 10px !important;
}

[data-testid="stExpander"] {
    background: var(--color-surface) !important;
    border: 1px solid var(--color-border) !important;
    border-radius: 0 !important;
}

[data-testid="stExpander"] summary {
    color: var(--color-text-muted) !important;
    font-family: var(--font-body) !important;
    font-size: 13px !important;
}

[data-testid="stForm"] {
    background: var(--color-surface) !important;
    border: 1px solid var(--color-border) !important;
    border-radius: 0 !important;
    padding: 1.5rem !important;
}

.stAlert {
    background: var(--color-surface) !important;
    border: 1px solid var(--color-border) !important;
    border-radius: 0 !important;
    font-family: var(--font-body) !important;
    font-size: 13px !important;
}

code {
    background: var(--color-purple-dim) !important;
    color: var(--color-purple-light) !important;
    border-radius: 2px !important;
    padding: 1px 5px !important;
    font-family: var(--font-mono) !important;
    font-size: 11px !important;
}

pre code { background: transparent !important; padding: 0 !important; }

pre {
    background: var(--color-surface) !important;
    border: 1px solid var(--color-border) !important;
    border-radius: 0 !important;
    padding: 1rem !important;
}

.stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {
    font-family: var(--font-display) !important;
    color: var(--color-text) !important;
    font-weight: 400 !important;
    letter-spacing: -0.01em !important;
}

.stMarkdown p, .stMarkdown li {
    font-family: var(--font-body) !important;
    color: var(--color-text-muted) !important;
    font-size: 14px !important;
    line-height: 1.75 !important;
}

/* File uploader */
[data-testid="stFileUploader"],
[data-testid="stFileUploaderDropzone"] {
    background: var(--color-surface) !important;
    border: 1px dashed var(--color-border-subtle) !important;
    border-radius: 0 !important;
}

[data-testid="stFileUploader"]:hover,
[data-testid="stFileUploaderDropzone"]:hover {
    border-color: var(--color-purple) !important;
}

[data-testid="stFileUploaderDropzone"] button {
    background: var(--color-purple-dim) !important;
    color: var(--color-purple-light) !important;
    border: 1px solid var(--color-purple) !important;
    border-radius: 0 !important;
    font-family: var(--font-mono) !important;
    font-size: 10px !important;
    letter-spacing: 0.08em !important;
}

/* Scrollbar */
::-webkit-scrollbar { width: 4px; height: 4px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: var(--color-border-subtle); border-radius: 2px; }
::-webkit-scrollbar-thumb:hover { background: var(--color-purple); }

/* Spinner */
.stSpinner > div { border-top-color: var(--color-purple) !important; }

/* Sidebar */
section[data-testid="stSidebar"],
[data-testid="stSidebarContent"],
[data-testid="stSidebarCollapsedControl"] {
    background-color: var(--color-bg) !important;
    border-right: 1px solid var(--color-border) !important;
}

/* Chrome */
[data-testid="stToolbar"],
[data-testid="stDecoration"],
[data-testid="stStatusWidget"] { display: none !important; }
#MainMenu, footer, header { visibility: hidden !important; }

/* Quick actions */
.pb-action > button {
    background: transparent !important;
    border: 1px solid var(--color-border) !important;
    color: var(--color-text-faint) !important;
    font-family: var(--font-body) !important;
    font-size: 12px !important;
    font-weight: 400 !important;
    text-align: left !important;
    justify-content: flex-start !important;
    border-radius: 0 !important;
}

.pb-action > button:hover {
    border-color: var(--color-purple) !important;
    color: var(--color-text-muted) !important;
    background: var(--color-surface) !important;
}

/* Tags */
.pb-tag {
    display: inline-block;
    font-family: var(--font-mono);
    font-size: 9px;
    font-weight: 500;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    padding: 2px 7px;
    background: var(--color-purple-dim);
    color: var(--color-purple-light);
    border: 1px solid rgba(83,74,183,0.3);
    border-radius: 2px;
}

/* Active nav state */
.pb-nav-active > button {
    border-left: 2px solid var(--color-purple) !important;
    background: rgba(83,74,183,0.08) !important;
    color: var(--color-purple-light) !important;
}

/* Staggered fade-in */
@keyframes pbFadeUp {
    from { opacity: 0; transform: translateY(6px); }
    to   { opacity: 1; transform: translateY(0); }
}
.pb-f1 { animation: pbFadeUp 0.35s ease both 0.04s; }
.pb-f2 { animation: pbFadeUp 0.35s ease both 0.10s; }
.pb-f3 { animation: pbFadeUp 0.35s ease both 0.17s; }
.pb-f4 { animation: pbFadeUp 0.35s ease both 0.24s; }

@media (prefers-reduced-motion: reduce) {
    .pb-f1, .pb-f2, .pb-f3, .pb-f4 { animation: none; }
}
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# RAG
# ══════════════════════════════════════════════════════════════════════════════

@st.cache_resource(show_spinner="Lade Sprachmodell ...")
def _rag_model():
    try:
        from sentence_transformers import SentenceTransformer
        return SentenceTransformer("sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
    except Exception:
        return None


def rag_search(query, top_k=5):
    if not INDEX_PATH.exists() or not META_PATH.exists():
        return []
    try:
        import faiss, numpy as np
        model = _rag_model()
        if not model:
            return []
        meta = json.loads(META_PATH.read_text())
        idx  = faiss.read_index(str(INDEX_PATH))
        q    = model.encode([query], convert_to_numpy=True, normalize_embeddings=True).astype(np.float32)
        scores, indices = idx.search(q, top_k * 3)
        seen, out = set(), []
        for s, i in zip(scores[0], indices[0]):
            if i < 0: continue
            p = meta[i]["path"]
            if p in seen: continue
            seen.add(p)
            out.append((float(s), p, meta[i]["preview"]))
            if len(out) >= top_k: break
        return out
    except Exception:
        return []


def get_rag_context(query):
    hits = rag_search(query, top_k=5)
    if not hits:
        return ""
    parts = ["## RELEVANTE VAULT-EINTRÄGE\n"]
    for _, path, _ in hits:
        try:
            text = (VAULT / path).read_text(errors="ignore")[:3000]
            parts.append(f"### {path}\n{text}\n")
        except Exception:
            pass
    return "\n".join(parts)


def rag_status():
    if INDEX_PATH.exists() and META_PATH.exists():
        meta = json.loads(META_PATH.read_text())
        ts   = datetime.fromtimestamp(INDEX_PATH.stat().st_mtime).strftime("%d.%m %H:%M")
        return len(meta), ts
    return None, None


def rebuild_rag_bg():
    def _r():
        subprocess.run(
            [sys.executable, str(VAULT / "_agent" / "rag_index.py")],
            cwd=str(VAULT), capture_output=True,
        )
        _rag_model.clear()
    threading.Thread(target=_r, daemon=True).start()


# ══════════════════════════════════════════════════════════════════════════════
# Outlook (Kalender) + Gmail
# ══════════════════════════════════════════════════════════════════════════════

def _load_module(name, path):
    try:
        import importlib
        spec = importlib.util.spec_from_file_location(name, path)
        mod  = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)
        return mod
    except Exception:
        return None


_OL = _load_module("outlook_client", VAULT / "_agent" / "outlook_client.py")
_GM = _load_module("gmail_client",   VAULT / "_agent" / "gmail_client.py")

outlook_cal_ok = bool(_OL and _OL.is_authenticated())
gmail_ok       = bool(_GM and _GM.is_authenticated())

# ── HTML-Frontend API-Server (Port 3001) ──────────────────────────────────────
def _start_api_server():
    """Startet brain_server.py im Hintergrund, damit die HTML-UI echte Daten bekommt."""
    import socket
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        if s.connect_ex(("localhost", 3001)) == 0:
            return  # läuft bereits
    server_path = VAULT / "_agent" / "brain_server.py"
    if server_path.exists():
        subprocess.Popen(
            [sys.executable, str(server_path)],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )

threading.Thread(target=_start_api_server, daemon=True).start()


def safe_get_emails(top=20, unread_only=False):
    if not _GM: return []
    try:    return _GM.get_emails(top=top, unread_only=unread_only)
    except: return []


def safe_get_events(days=7):
    if not _OL: return []
    try:    return _OL.get_calendar_events(days=days)
    except: return []


# ══════════════════════════════════════════════════════════════════════════════
# Claude Tools
# ══════════════════════════════════════════════════════════════════════════════

TOOLS = [
    {
        "name": "send_email",
        "description": "Sendet eine E-Mail über Sebastians Gmail.",
        "input_schema": {
            "type": "object",
            "properties": {
                "to":      {"type": "string"},
                "subject": {"type": "string"},
                "body":    {"type": "string"},
                "cc":      {"type": "string"},
            },
            "required": ["to", "subject", "body"],
        },
    },
    {
        "name": "reply_to_email",
        "description": "Antwortet auf eine bestehende Gmail-E-Mail.",
        "input_schema": {
            "type": "object",
            "properties": {
                "message_id":      {"type": "string"},
                "thread_id":       {"type": "string"},
                "to":              {"type": "string"},
                "subject":         {"type": "string"},
                "orig_message_id": {"type": "string"},
                "orig_references": {"type": "string"},
                "body":            {"type": "string"},
            },
            "required": ["message_id", "thread_id", "to", "body"],
        },
    },
    {
        "name": "create_calendar_event",
        "description": "Erstellt einen Termin im Outlook-Kalender.",
        "input_schema": {
            "type": "object",
            "properties": {
                "subject":   {"type": "string"},
                "start_dt":  {"type": "string"},
                "end_dt":    {"type": "string"},
                "body":      {"type": "string"},
                "attendees": {"type": "string"},
                "location":  {"type": "string"},
            },
            "required": ["subject", "start_dt", "end_dt"],
        },
    },
    {
        "name": "get_inbox_emails",
        "description": "Liest E-Mails aus dem Posteingang.",
        "input_schema": {
            "type": "object",
            "properties": {
                "top":         {"type": "integer"},
                "unread_only": {"type": "boolean"},
            },
        },
    },
    {
        "name": "get_upcoming_events",
        "description": "Ruft bevorstehende Kalendertermine ab.",
        "input_schema": {
            "type": "object",
            "properties": {
                "days": {"type": "integer"},
            },
        },
    },
    {
        "name": "save_to_memory",
        "description": (
            "Speichert eine Regel, Korrektur, Erkenntnis oder Prozessbeschreibung dauerhaft im Brain-Gedächtnis. "
            "Nutze dieses Tool IMMER wenn Sebastian dich korrigiert, etwas erklärt, 'merke dir' oder 'lerne' sagt, "
            "oder wenn du erkennst dass eine Information für zukünftige Gespräche wichtig ist."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "typ": {
                    "type": "string",
                    "enum": ["REGEL", "KORREKTUR", "PROZESS", "KONTEXT"],
                    "description": "REGEL=Ablaufregel, KORREKTUR=Fehlerkorrektur, PROZESS=Arbeitsablauf, KONTEXT=Hintergrundwissen",
                },
                "inhalt": {
                    "type": "string",
                    "description": "Der Inhalt der zu merkenden Information, präzise formuliert.",
                },
            },
            "required": ["typ", "inhalt"],
        },
    },
    {
        "name": "update_context",
        "description": (
            "Aktualisiert den laufenden Arbeitskontext: Aufgaben hinzufügen/abhaken, Notizen. "
            "Nutze dieses Tool wenn neue Aufgaben entstehen, erledigt werden, oder Sebastian den Stand mitteilt."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "aktion": {
                    "type": "string",
                    "enum": ["aufgabe_hinzufügen", "aufgabe_erledigt", "priorität_ändern", "notiz_hinzufügen"],
                },
                "inhalt": {
                    "type": "string",
                    "description": "Text der Aufgabe oder Notiz.",
                },
            },
            "required": ["aktion", "inhalt"],
        },
    },
    {
        "name": "vault_operation",
        "description": (
            "Führt eine Strukturoperation im Vault durch: Ordner erstellen (mit Vorlage), umbenennen, "
            "Datei verschieben, Ordnerinhalt auflisten. "
            "Nutze dieses Tool wenn Sebastian die Vault-Struktur ändern will oder einen neuen Ordner braucht."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "aktion": {
                    "type": "string",
                    "enum": ["ordner_erstellen", "ordner_umbenennen", "datei_verschieben", "ordner_auflisten"],
                },
                "pfad": {
                    "type": "string",
                    "description": "Relativer Pfad vom Vault-Root (z.B. 'Kunden/NeueFirma')",
                },
                "neuer_pfad": {
                    "type": "string",
                    "description": "Neuer Pfad (nur für umbenennen und verschieben)",
                },
                "vorlage": {
                    "type": "string",
                    "enum": ["kunde", "lead", "projekt", "leer"],
                    "description": "Ordnervorlage mit Standard-README (nur für ordner_erstellen)",
                },
            },
            "required": ["aktion", "pfad"],
        },
    },
    {
        "name": "update_prozessia_profile",
        "description": (
            "Aktualisiert das Prozessia-Masterprofil (prozessia.md). "
            "Nutze dieses Tool wenn sich Kunden, Prioritäten, Produkte, Team oder Kerninformationen ändern."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "abschnitt": {
                    "type": "string",
                    "description": "Name des Abschnitts (z.B. 'Aktive Kunden', 'Offene Prioritäten', 'Produkte')",
                },
                "inhalt": {
                    "type": "string",
                    "description": "Neuer Inhalt oder Zusatz",
                },
                "aktion": {
                    "type": "string",
                    "enum": ["anhängen", "ersetzen"],
                    "description": "anhängen = füge zu bestehendem Inhalt hinzu, ersetzen = überschreibe den Abschnitt",
                },
            },
            "required": ["abschnitt", "inhalt", "aktion"],
        },
    },
]


def execute_tool(name, inp):
    try:
        if name == "send_email":
            if not _GM: return "Gmail nicht verbunden."
            return _GM.send_email(inp["to"], inp["subject"], inp["body"], inp.get("cc"))
        elif name == "reply_to_email":
            if not _GM: return "Gmail nicht verbunden."
            return _GM.reply_email(
                inp["message_id"], inp["thread_id"], inp["to"],
                inp.get("subject", ""), inp.get("orig_message_id", inp["message_id"]),
                inp.get("orig_references", ""), inp["body"],
            )
        elif name == "create_calendar_event":
            if not _OL: return "Outlook Kalender nicht verbunden."
            att = [a.strip() for a in inp["attendees"].split(",")] if inp.get("attendees") else None
            return _OL.create_calendar_event(
                inp["subject"], inp["start_dt"], inp["end_dt"],
                inp.get("body", ""), att, inp.get("location", ""),
            )
        elif name == "get_inbox_emails":
            if not _GM: return "Gmail nicht verbunden."
            emails = _GM.get_emails(top=inp.get("top", 10), unread_only=inp.get("unread_only", False))
            return json.dumps([{
                "id":      e["id"],
                "from":    e["from"]["emailAddress"]["address"],
                "name":    e["from"]["emailAddress"].get("name", ""),
                "subject": e["subject"],
                "preview": e.get("bodyPreview", "")[:200],
                "date":    e["receivedDateTime"][:16].replace("T", " "),
                "unread":  not e.get("isRead", True),
            } for e in emails], ensure_ascii=False)
        elif name == "get_upcoming_events":
            if not _OL: return "Outlook Kalender nicht verbunden."
            events = _OL.get_calendar_events(days=inp.get("days", 7))
            return json.dumps([{
                "subject":  ev["subject"],
                "start":    ev["start"]["dateTime"][:16].replace("T", " "),
                "end":      ev["end"]["dateTime"][11:16],
                "location": ev.get("location", {}).get("displayName", ""),
                "preview":  ev.get("bodyPreview", "")[:100],
            } for ev in events], ensure_ascii=False)
        elif name == "save_to_memory":
            datum  = datetime.now().strftime("%Y-%m-%d %H:%M")
            typ    = inp.get("typ", "KONTEXT")
            inhalt = inp.get("inhalt", "")
            entry  = f"- [{datum}] {inhalt}\n"

            text = MEMORY_PATH.read_text() if MEMORY_PATH.exists() else "# Brain-Gedächtnis\n\n"
            header = f"\n## {typ}\n"

            if f"## {typ}" in text:
                text = text.replace(f"## {typ}\n", f"## {typ}\n{entry}")
            else:
                text = text.rstrip() + f"\n\n## {typ}\n{entry}"

            MEMORY_PATH.write_text(text)
            load_vault_context.clear()
            return f"Gespeichert: [{typ}] {inhalt}"

        elif name == "update_context":
            datum  = datetime.now().strftime("%Y-%m-%d %H:%M")
            aktion = inp.get("aktion", "notiz_hinzufügen")
            inhalt = inp.get("inhalt", "")

            if CONTEXT_PATH.exists():
                text = CONTEXT_PATH.read_text()
            else:
                text = f"# Aktueller Kontext\n\nAktualisiert: {datum}\n\n## Offene Aufgaben\n\n## Notizen\n"

            if aktion == "aufgabe_hinzufügen":
                if "## Offene Aufgaben\n" in text:
                    text = text.replace("## Offene Aufgaben\n", f"## Offene Aufgaben\n- [ ] {inhalt}\n")
                else:
                    text += f"\n## Offene Aufgaben\n- [ ] {inhalt}\n"
            elif aktion == "aufgabe_erledigt":
                text = text.replace(f"- [ ] {inhalt}", f"- [x] {inhalt}")
            else:
                entry = f"- [{datum}] {inhalt}\n"
                if "## Notizen\n" in text:
                    text = text.replace("## Notizen\n", f"## Notizen\n{entry}")
                else:
                    text += f"\n## Notizen\n{entry}"

            text = re.sub(r"Aktualisiert: .+", f"Aktualisiert: {datum}", text)
            CONTEXT_PATH.write_text(text)
            load_vault_context.clear()
            return f"Kontext aktualisiert: {aktion} — {inhalt}"

        elif name == "vault_operation":
            aktion   = inp.get("aktion", "")
            pfad_rel = inp.get("pfad", "").lstrip("/")
            neu_rel  = inp.get("neuer_pfad", "").lstrip("/")
            abs_pfad = (VAULT / pfad_rel).resolve()

            try:
                abs_pfad.relative_to(VAULT.resolve())
            except ValueError:
                return "Fehler: Pfad außerhalb des Vaults nicht erlaubt."

            if aktion == "ordner_erstellen":
                abs_pfad.mkdir(parents=True, exist_ok=True)
                vorlage = inp.get("vorlage", "leer")
                if vorlage in ("kunde", "lead"):
                    tpl = f"# {abs_pfad.name}\n\nStatus: {'Aktiv' if vorlage=='kunde' else 'Lead'}\n\n## Kontakt\n\n## Projekte\n\n## Notizen\n"
                    (abs_pfad / "README.md").write_text(tpl)
                load_vault_context.clear()
                return f"Ordner erstellt: {pfad_rel}"

            elif aktion == "ordner_umbenennen":
                if abs_pfad.exists() and neu_rel:
                    (VAULT / neu_rel).parent.mkdir(parents=True, exist_ok=True)
                    abs_pfad.rename((VAULT / neu_rel).resolve())
                    load_vault_context.clear()
                    return f"Umbenannt: {pfad_rel} → {neu_rel}"
                return f"Fehler: {pfad_rel} nicht gefunden."

            elif aktion == "datei_verschieben":
                if abs_pfad.exists() and neu_rel:
                    abs_neu = (VAULT / neu_rel).resolve()
                    abs_neu.parent.mkdir(parents=True, exist_ok=True)
                    shutil.move(str(abs_pfad), str(abs_neu))
                    load_vault_context.clear()
                    return f"Verschoben: {pfad_rel} → {neu_rel}"
                return f"Fehler: {pfad_rel} nicht gefunden."

            elif aktion == "ordner_auflisten":
                if abs_pfad.exists():
                    items = sorted(
                        f.name + ("/" if f.is_dir() else "")
                        for f in abs_pfad.iterdir()
                        if not f.name.startswith(".")
                    )
                    return f"Inhalt von {pfad_rel} ({len(items)} Einträge):\n" + "\n".join(f"  {i}" for i in items[:60])
                return f"Nicht gefunden: {pfad_rel}"

            return f"Unbekannte Aktion: {aktion}"

        elif name == "update_prozessia_profile":
            abschnitt = inp.get("abschnitt", "")
            inhalt    = inp.get("inhalt", "")
            aktion    = inp.get("aktion", "anhängen")

            if not PROZESSIA_PATH.exists():
                return "prozessia.md nicht gefunden."

            text   = PROZESSIA_PATH.read_text()
            header = f"## {abschnitt}"

            if header in text:
                idx   = text.index(header) + len(header)
                rest  = text[idx:]
                nxt   = rest.find("\n## ")
                if aktion == "anhängen":
                    entry = f"\n- {inhalt}"
                    text  = text[:idx] + (rest[:nxt] + entry + rest[nxt:] if nxt >= 0 else rest.rstrip() + entry + "\n")
                elif aktion == "ersetzen":
                    neue  = f"\n{inhalt}\n"
                    text  = text[:idx] + (neue + rest[nxt:] if nxt >= 0 else neue)
            else:
                text = text.rstrip() + f"\n\n{header}\n- {inhalt}\n"

            PROZESSIA_PATH.write_text(text)
            load_vault_context.clear()
            return f"Prozessia-Profil aktualisiert: '{abschnitt}'"

        return f"Unbekanntes Tool: {name}"
    except Exception as e:
        return f"Fehler: {e}"


# ══════════════════════════════════════════════════════════════════════════════
# Vault-Hilfsfunktionen
# ══════════════════════════════════════════════════════════════════════════════

@st.cache_data(ttl=120)
def load_vault_context():
    parts = []

    if PROZESSIA_PATH.exists():
        parts.append(f"## MASTERPROFIL PROZESSIA\n{PROZESSIA_PATH.read_text()}\n")

    if MEMORY_PATH.exists():
        mem = MEMORY_PATH.read_text().strip()
        if mem:
            parts.append(f"## BRAIN-GEDÄCHTNIS (gelernte Regeln & Korrekturen)\n{mem}\n")

    if CONTEXT_PATH.exists():
        ctx = CONTEXT_PATH.read_text().strip()
        if ctx:
            parts.append(f"## AKTUELLER KONTEXT & AUFGABEN\n{ctx}\n")

    kunden_dir = VAULT / "Kunden"
    if kunden_dir.exists():
        for md in sorted(kunden_dir.rglob("*.md"))[:12]:
            try:
                parts.append(f"## KUNDE: {md.parent.name}\n{md.read_text()[:2000]}\n")
            except Exception:
                pass

    today = VAULT / "_agent" / "daily" / f"{datetime.now().strftime('%Y-%m-%d')}.md"
    if today.exists():
        parts.append(f"## HEUTE\n{today.read_text()}\n")

    return "\n".join(parts)


def keyword_search(query):
    kws     = query.lower().split()
    results = []
    for md in VAULT.rglob("*.md"):
        if any(x in str(md) for x in ["_agent", "_inbox", ".git"]):
            continue
        try:
            c = md.read_text(errors="ignore")
            if all(k in c.lower() for k in kws):
                results.append((str(md.relative_to(VAULT)), c[:250].replace("\n", " ")))
        except Exception:
            pass
    return results[:8]


def list_vault_files():
    files = []
    for f in VAULT.rglob("*"):
        if f.is_file() and not any(x in str(f) for x in ["_inbox", ".git", ".obsidian/plugins", "_agent/logs"]):
            files.append(str(f.relative_to(VAULT)))
    return sorted(files)


def extract_text(raw: bytes, filename: str):
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".pdf":
            import PyPDF2
            return " ".join(p.extract_text() or "" for p in PyPDF2.PdfReader(io.BytesIO(raw)).pages)[:8000]
        if ext == ".docx":
            from docx import Document
            return " ".join(p.text for p in Document(io.BytesIO(raw)).paragraphs)[:8000]
        if ext == ".xlsx":
            import openpyxl
            wb   = openpyxl.load_workbook(io.BytesIO(raw), read_only=True)
            rows = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    rows.append(" ".join(str(c) for c in row if c))
            return " ".join(rows)[:8000]
        if ext in {".txt", ".md", ".csv", ".json", ".html"}:
            return raw.decode("utf-8", errors="ignore")[:8000]
    except Exception:
        pass
    return None


def fmt_email_time(dt_str):
    try:
        dt   = datetime.fromisoformat(dt_str.replace("Z", "+00:00"))
        now  = datetime.now(dt.tzinfo)
        diff = (now - dt).days
        if diff == 0:  return dt.strftime("%H:%M")
        if diff < 7:   return dt.strftime("%a")
        return dt.strftime("%d.%m.")
    except Exception:
        return dt_str[:10]


def strip_html(html):
    return re.sub(r"<[^>]+>", " ", html or "").strip()


# ══════════════════════════════════════════════════════════════════════════════
# Agent
# ══════════════════════════════════════════════════════════════════════════════

def run_agent(client, messages, system, tools, placeholder):
    full_response = ""
    tools_active  = (gmail_ok or outlook_cal_ok) and st.session_state.get("use_tools", True)
    active_tools  = tools if tools_active else []

    while True:
        with client.messages.stream(
            model="claude-sonnet-4-6",
            max_tokens=2048,
            system=system,
            messages=messages,
            tools=active_tools,
        ) as stream:
            for text in stream.text_stream:
                full_response += text
                placeholder.markdown(full_response + "▌")
            final = stream.get_final_message()

        if final.stop_reason != "tool_use":
            break

        assistant_content = []
        tool_results      = []
        for block in final.content:
            if block.type == "text":
                assistant_content.append({"type": "text", "text": block.text})
            elif block.type == "tool_use":
                assistant_content.append({
                    "type": "tool_use", "id": block.id,
                    "name": block.name, "input": block.input,
                })
                with st.spinner(f"{block.name} ..."):
                    result = execute_tool(block.name, block.input)
                tool_results.append({
                    "type": "tool_result",
                    "tool_use_id": block.id,
                    "content": str(result),
                })

        messages = messages + [
            {"role": "assistant", "content": assistant_content},
            {"role": "user",      "content": tool_results},
        ]

    placeholder.markdown(full_response)
    return full_response


# ══════════════════════════════════════════════════════════════════════════════
# Session State
# ══════════════════════════════════════════════════════════════════════════════

defaults = {
    "messages":       [],
    "uploaded_docs":  {},
    "rag_enabled":    True,
    "use_tools":      True,
    "emails":         None,
    "emails_ts":      None,
    "events":         None,
    "events_ts":      None,
    "selected_email": None,
    "compose_open":   False,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ══════════════════════════════════════════════════════════════════════════════
# Header
# ══════════════════════════════════════════════════════════════════════════════

rag_chunks, rag_ts = rag_status()

rag_chip = (
    f'<span class="pb-chip chip-ok">RAG &thinsp; {rag_chunks}</span>'
    if rag_chunks else
    '<span class="pb-chip chip-off">RAG</span>'
)
gm_chip = (
    '<span class="pb-chip chip-on">Gmail</span>'
    if gmail_ok else
    '<span class="pb-chip chip-off">Gmail</span>'
)
ol_chip = (
    '<span class="pb-chip chip-on">Kalender</span>'
    if outlook_cal_ok else
    '<span class="pb-chip chip-off">Kalender</span>'
)

today_hdr = datetime.now().strftime("%A, %-d. %B %Y")

STARS = [
    (6,45),(18,90),(32,38),(9,140),(44,72),(22,175),(7,225),
    (38,115),(14,265),(28,198),(50,52),(11,308),(40,150),(24,335),
    (5,380),(47,255),(17,412),(33,82),(52,345),(13,445),(42,195),
    (26,470),(9,502),(36,535),(21,568),(51,500),(16,600),
]
stars_html = "".join(
    f'<span class="pb-star" style="top:{t}px;right:{r}px"></span>'
    for t, r in STARS
)

st.markdown(f"""
<div class="pb-header pb-f1">
    {stars_html}
    <div>
        <div class="pb-wordmark">Prozessia <span class="accent">Brain</span></div>
        <div class="pb-meta">{today_hdr}</div>
    </div>
    <div class="pb-chips">{rag_chip}{gm_chip}{ol_chip}</div>
</div>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# Tabs
# ══════════════════════════════════════════════════════════════════════════════

tab_chat, tab_mail, tab_cal = st.tabs(["Chat", "Postfach", "Kalender"])


# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 – CHAT
# ══════════════════════════════════════════════════════════════════════════════

with tab_chat:
    col_chat, col_side = st.columns([3, 1], gap="large")

    with col_side:

        if rag_chunks:
            st.toggle("RAG aktiv", key="rag_enabled")
            st.caption(f"Stand {rag_ts} · {rag_chunks} Dokumente")
        else:
            st.markdown('<div class="pb-label">Wissensindex</div>', unsafe_allow_html=True)
            if st.button("Index aufbauen", use_container_width=True):
                with st.spinner("Indiziere Vault ..."):
                    r = subprocess.run(
                        [sys.executable, str(VAULT / "_agent" / "rag_index.py")],
                        cwd=str(VAULT), capture_output=True, text=True, timeout=300,
                    )
                    st.rerun() if r.returncode == 0 else st.error(r.stderr[-300:])

        if gmail_ok or outlook_cal_ok:
            st.toggle("Tools aktiv", key="use_tools")

        st.divider()

        # Upload
        st.markdown('<div class="pb-label">Dokument</div>', unsafe_allow_html=True)
        uploaded = st.file_uploader(
            "Upload",
            type=["pdf", "docx", "xlsx", "pptx", "txt", "md", "csv", "json"],
            accept_multiple_files=True,
            label_visibility="collapsed",
        )
        if uploaded:
            for f in uploaded:
                if f.name not in st.session_state.uploaded_docs:
                    raw = f.read()
                    INBOX.mkdir(parents=True, exist_ok=True)
                    (INBOX / f.name).write_bytes(raw)
                    threading.Thread(
                        target=lambda: subprocess.run(
                            [sys.executable, str(VAULT / "_agent" / "heartbeat.py")],
                            cwd=str(VAULT), capture_output=True,
                            env={**os.environ, "ANTHROPIC_API_KEY": API_KEY or ""},
                        ), daemon=True,
                    ).start()
                    text = extract_text(raw, f.name)
                    if text:
                        st.session_state.uploaded_docs[f.name] = text
                    st.caption(f.name)

        # Vault-Datei öffnen
        st.markdown('<div class="pb-label">Vault</div>', unsafe_allow_html=True)
        vault_files = list_vault_files()
        sel = st.selectbox("Datei", ["— auswählen —"] + vault_files, label_visibility="collapsed")
        if sel != "— auswählen —" and sel not in st.session_state.uploaded_docs:
            fp  = VAULT / sel
            ext = Path(sel).suffix.lower()
            try:
                if ext in {".md", ".txt", ".csv", ".json"}:
                    content = fp.read_text(errors="ignore")[:8000]
                elif ext == ".pdf":
                    import PyPDF2
                    content = " ".join(p.extract_text() or "" for p in PyPDF2.PdfReader(str(fp)).pages)[:8000]
                elif ext == ".docx":
                    from docx import Document
                    content = " ".join(p.text for p in Document(fp).paragraphs)[:8000]
                elif ext == ".xlsx":
                    import openpyxl
                    wb   = openpyxl.load_workbook(fp, read_only=True)
                    rows = []
                    for ws in wb.worksheets:
                        for row in ws.iter_rows(values_only=True):
                            rows.append(" ".join(str(c) for c in row if c))
                    content = " ".join(rows)[:8000]
                else:
                    content = None
                if content:
                    st.session_state.uploaded_docs[sel] = content
                    st.caption("Geladen")
            except Exception as e:
                st.error(str(e))

        # Aktive Dokumente
        if st.session_state.uploaded_docs:
            st.markdown('<div class="pb-label">Im Kontext</div>', unsafe_allow_html=True)
            for name in list(st.session_state.uploaded_docs):
                c1, c2 = st.columns([5, 1])
                c1.caption(Path(name).name)
                if c2.button("×", key=f"rm_{name}"):
                    del st.session_state.uploaded_docs[name]
                    st.rerun()

        st.divider()

        # Schnellzugriff
        st.markdown('<div class="pb-label">Schnellzugriff</div>', unsafe_allow_html=True)
        quick = {
            "Offene Aufgaben":   "Was sind alle offenen Aufgaben und Prioritäten?",
            "Aktive Kunden":     "Übersicht aller aktiven Kunden und Warm Leads.",
            "Termine":           "Welche Termine stehen als nächstes an?",
            "Finanzstatus":      "Fasse den finanziellen Stand zusammen.",
            "Posteingang":       "Lies meine letzten 10 E-Mails und fasse sie kurz zusammen.",
        }
        for label, prompt in quick.items():
            st.markdown('<div class="pb-action">', unsafe_allow_html=True)
            if st.button(label, use_container_width=True, key=f"q_{label}"):
                st.session_state.messages.append({"role": "user", "content": prompt})
                st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)

        st.divider()

        # Suche
        st.markdown('<div class="pb-label">Suche</div>', unsafe_allow_html=True)
        sq = st.text_input("Suche", label_visibility="collapsed", placeholder="Schaufler, NDA, Angebot ...")
        if sq:
            hits = rag_search(sq, top_k=5) if rag_chunks else []
            if hits:
                for _s, path, preview in hits:
                    st.markdown(f"**{Path(path).name}**")
                    st.caption(preview[:90] + "...")
            else:
                kw = keyword_search(sq)
                for path, snippet in kw:
                    st.markdown(f"**{Path(path).name}**")
                    st.caption(snippet[:90] + "...")
                if not kw:
                    st.caption("Keine Treffer.")

        st.divider()

        c1, c2 = st.columns(2)
        if c1.button("Chat leeren",  use_container_width=True): st.session_state.messages = []; st.rerun()
        if c2.button("Neu laden",    use_container_width=True): st.cache_data.clear(); st.rerun()
        if st.button("RAG neu aufbauen", use_container_width=True): rebuild_rag_bg(); st.caption("Läuft im Hintergrund ...")

    with col_chat:
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        prompt = st.chat_input("Frag dein Second Brain ...")

        if prompt:
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                vault_ctx = load_vault_context()
                rag_ctx   = get_rag_context(prompt) if (st.session_state.rag_enabled and rag_chunks) else ""
                doc_ctx   = ""
                if st.session_state.uploaded_docs:
                    doc_ctx = "\n\n## GEÖFFNETE DOKUMENTE\n"
                    for name, content in st.session_state.uploaded_docs.items():
                        doc_ctx += f"\n### {name}\n{content}\n"

                today_str = datetime.now().strftime("%A, %d. %B %Y, %H:%M")
                system = f"""Du bist das Second Brain von Sebastian Spuhler (Prozessia GbR, Saarbrücken).
Heute: {today_str}
Antworte auf Deutsch, präzise und direkt. Nutze Markdown.

LERNFÄHIGKEIT — Nutze diese Tools proaktiv ohne explizit gefragt zu werden:
• save_to_memory → wenn Sebastian dich korrigiert, etwas erklärt, "merke dir" / "vergiss nicht" / "lerne" sagt
• update_context → wenn neue Aufgaben entstehen, erledigt werden oder sich Prioritäten ändern
• vault_operation → wenn Sebastian Ordner erstellen, umbenennen oder Dateien verschieben will
• update_prozessia_profile → wenn sich Kunden, Prioritäten, Produkte oder Team-Infos ändern

WEITERE TOOLS: E-Mails senden, Kalendertermine anlegen – nutze sie wenn Sebastian das möchte.
Wenn du eine E-Mail schreiben sollst: formuliere sie professionell auf Deutsch, dann sende sie direkt.

{vault_ctx}
{rag_ctx}
{doc_ctx}

Wenn du etwas nicht weißt, sag es direkt."""

                history = [
                    {"role": m["role"], "content": m["content"]}
                    for m in st.session_state.messages[-12:]
                ]
                client      = anthropic.Anthropic(api_key=API_KEY)
                placeholder = st.empty()
                response    = run_agent(client, history, system, TOOLS, placeholder)
                st.session_state.messages.append({"role": "assistant", "content": response})


# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 – POSTFACH
# ══════════════════════════════════════════════════════════════════════════════

with tab_mail:
    if not gmail_ok:
        st.markdown("### Gmail nicht verbunden")
        st.markdown("""
Login ausführen:
```bash
python3 _agent/gmail_setup.py
```
Danach UI neu starten.
        """)
        st.stop()

    cm1, cm2, cm3, _ = st.columns([1, 1, 1, 3])
    if cm1.button("Aktualisieren",  use_container_width=True): st.session_state.emails = None
    if cm2.button("Neue Mail",      use_container_width=True): st.session_state.compose_open = not st.session_state.compose_open
    unread_filter = cm3.checkbox("Nur ungelesen")

    if st.session_state.compose_open:
        with st.form("compose_form"):
            st.markdown("**Neue E-Mail**")
            to_addr = st.text_input("An")
            cc_addr = st.text_input("CC (optional)")
            subject = st.text_input("Betreff")
            body    = st.text_area("Text", height=140)
            c1, _ = st.columns([1, 4])
            if c1.form_submit_button("Senden", type="primary"):
                if to_addr and subject:
                    st.success(_GM.send_email(to_addr, subject, body, cc_addr or None))
                    st.session_state.compose_open = False
                    st.rerun()
                else:
                    st.error("An und Betreff sind Pflicht.")

    if st.session_state.emails is None:
        with st.spinner("Lade E-Mails ..."):
            st.session_state.emails    = safe_get_emails(top=25, unread_only=unread_filter)
            st.session_state.emails_ts = datetime.now()

    emails = st.session_state.emails or []

    if not emails:
        st.markdown('<p style="color:#3a3a4a;font-size:13px;margin-top:1.5rem;">Keine E-Mails gefunden.</p>', unsafe_allow_html=True)
    else:
        ts_str = st.session_state.emails_ts.strftime("%H:%M") if st.session_state.emails_ts else ""
        st.markdown(
            f'<p style="color:#3a3a4a;font-size:10px;letter-spacing:0.08em;text-transform:uppercase;'
            f'margin-bottom:14px;">{len(emails)} Nachrichten &nbsp;·&nbsp; {ts_str}</p>',
            unsafe_allow_html=True,
        )

        col_list, col_detail = st.columns([5, 7], gap="large")

        with col_list:
            for i, email in enumerate(emails):
                unread = not email.get("isRead", True)
                sender = email["from"]["emailAddress"].get("name") or email["from"]["emailAddress"]["address"]
                subj   = email.get("subject", "(kein Betreff)")
                t      = fmt_email_time(email["receivedDateTime"])

                dot      = '<span class="pb-unread-dot"></span>' if unread else ''
                subj_cls = "pb-card-subject unread-subj" if unread else "pb-card-subject"

                st.markdown(f"""
<div class="pb-card" style="cursor:pointer;">
    <div style="display:flex;justify-content:space-between;align-items:flex-start;">
        <span class="pb-card-sender">{dot}{sender[:30]}</span>
        <span class="pb-card-time">{t}</span>
    </div>
    <div class="{subj_cls}">{subj[:55]}</div>
</div>
""", unsafe_allow_html=True)
                if st.button("Öffnen", key=f"email_{i}", use_container_width=True):
                    st.session_state.selected_email = email

        with col_detail:
            sel = st.session_state.selected_email
            if sel:
                st.markdown(f"### {sel.get('subject', '(kein Betreff)')}")
                c1, c2 = st.columns(2)
                addr   = sel['from']['emailAddress']
                c1.caption(f"Von: {addr.get('name', '')}  ·  {addr['address']}")
                c2.caption(f"Datum: {sel['receivedDateTime'][:16].replace('T', ' ')}")
                st.divider()
                body_html = sel.get("body", {}).get("content", sel.get("bodyPreview", ""))
                body_text = strip_html(body_html) if sel.get("body", {}).get("contentType") == "html" else body_html
                st.text_area("", value=body_text[:3000], height=280, disabled=True, label_visibility="collapsed")
                st.divider()
                with st.expander("Antworten"):
                    with st.form(f"reply_{sel['id'][:8]}"):
                        reply_body = st.text_area("Antwort", height=120)
                        if st.form_submit_button("Absenden", type="primary"):
                            st.success(_GM.reply_email(
                                message_id=sel["id"], thread_id=sel["threadId"],
                                to=sel["from"], orig_subject=sel["subject"],
                                orig_message_id=sel.get("message_id", sel["id"]),
                                orig_references=sel.get("references", ""),
                                body=reply_body,
                            ))
            else:
                st.markdown(
                    '<p style="color:#3a3a4a;font-size:13px;margin-top:2rem;">Nachricht aus der Liste auswählen.</p>',
                    unsafe_allow_html=True,
                )


# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 – KALENDER
# ══════════════════════════════════════════════════════════════════════════════

with tab_cal:
    if not outlook_cal_ok:
        st.markdown("### Outlook Kalender nicht verbunden")
        st.markdown("""
Login ausführen:
```bash
python3 _agent/ms_login.py
```
Danach UI neu starten.
        """)
        st.stop()

    cc1, cc2, _ = st.columns([1, 1, 4])
    if cc1.button("Aktualisieren", use_container_width=True): st.session_state.events = None
    days_ahead = cc2.selectbox(
        "Zeitraum", [7, 14, 30],
        format_func=lambda d: f"{d} Tage",
        label_visibility="collapsed",
    )

    if st.session_state.events is None:
        with st.spinner("Lade Kalender ..."):
            st.session_state.events    = safe_get_events(days=days_ahead)
            st.session_state.events_ts = datetime.now()

    events = st.session_state.events or []
    col_events, col_form = st.columns([3, 2], gap="large")

    with col_events:
        if not events:
            st.markdown(
                f'<p style="color:#3a3a4a;font-size:13px;">Keine Termine in den nächsten {days_ahead} Tagen.</p>',
                unsafe_allow_html=True,
            )
        else:
            by_day   = defaultdict(list)
            for ev in events:
                by_day[ev["start"]["dateTime"][:10]].append(ev)

            today    = datetime.now().date()
            wd_short = ["Mo", "Di", "Mi", "Do", "Fr", "Sa", "So"]

            for day_str in sorted(by_day.keys()):
                day_dt = datetime.fromisoformat(day_str).date()
                diff   = (day_dt - today).days
                if diff == 0:   label = "Heute"
                elif diff == 1: label = "Morgen"
                else:           label = f"{wd_short[day_dt.weekday()]} · {day_dt.strftime('%-d. %b')}"

                st.markdown(f'<div class="pb-day">{label}</div>', unsafe_allow_html=True)

                for ev in by_day[day_str]:
                    s     = ev["start"]["dateTime"][11:16]
                    e     = ev["end"]["dateTime"][11:16]
                    loc   = ev.get("location", {}).get("displayName", "")
                    subj  = ev.get("subject", "")
                    prev  = ev.get("bodyPreview", "")[:80]
                    meta  = " &nbsp;·&nbsp; ".join(filter(None, [loc, prev]))

                    st.markdown(f"""
<div class="pb-event">
    <div class="pb-event-time">{s} – {e}</div>
    <div class="pb-event-title">{subj}</div>
    {"<div class='pb-event-meta'>" + meta + "</div>" if meta else ""}
</div>
""", unsafe_allow_html=True)

    with col_form:
        st.markdown('<div class="pb-label">Neuer Termin</div>', unsafe_allow_html=True)
        with st.form("new_event"):
            ev_title = st.text_input("Titel")
            ev_date  = st.date_input("Datum", value=datetime.now().date())
            c1, c2   = st.columns(2)
            ev_start = c1.time_input("Von", value=time(10, 0))
            ev_end   = c2.time_input("Bis", value=time(11, 0))
            ev_loc   = st.text_input("Ort (optional)")
            ev_att   = st.text_input("Teilnehmer (kommagetrennt, optional)")
            ev_body  = st.text_area("Beschreibung (optional)", height=80)

            if st.form_submit_button("Termin anlegen", type="primary"):
                if ev_title:
                    att    = [a.strip() for a in ev_att.split(",") if a.strip()] if ev_att else None
                    result = _OL.create_calendar_event(
                        ev_title,
                        datetime.combine(ev_date, ev_start),
                        datetime.combine(ev_date, ev_end),
                        ev_body, att, ev_loc,
                    )
                    st.success(result)
                    st.session_state.events = None
                    st.rerun()
                else:
                    st.error("Titel ist Pflicht.")
